
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import json
import os
import tempfile
import re
from typing import Optional
import PyPDF2
from docx import Document
from openai import OpenAI

app = FastAPI()

app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

AI_API_KEY = os.environ.get("AI_API_KEY", "sk-在这里粘贴你的DeepSeek密钥")
AI_BASE_URL = "https://api.deepseek.com"
AI_MODEL = "deepseek-chat"

client = OpenAI(api_key=AI_API_KEY, base_url=AI_BASE_URL)


def clean_text(text):
    return text.encode('utf-8', errors='ignore').decode('utf-8')


def get_system_prompt():
    p = os.path.join(os.path.dirname(os.path.abspath(__file__)), "prompt.txt")
    if os.path.exists(p):
        f = open(p, "r", encoding="utf-8")
        txt = f.read()
        f.close()
        return txt
    return "你是报告解析专家。将报告解析为JSON。只输出JSON。"


def extract_pdf(path):
    text = ""
    f = open(path, "rb")
    reader = PyPDF2.PdfReader(f)
    for page in reader.pages:
        t = page.extract_text()
        if t:
            text += t + "\n"
    f.close()
    if not text.strip():
        raise Exception("PDF内容为空")
    return clean_text(text)


def extract_docx(path):
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                if cell.text.strip():
                    cells.append(cell.text.strip())
            if cells:
                lines.append(" | ".join(cells))
    text = "\n".join(lines)
    if not text.strip():
        raise Exception("DOCX内容为空")
    return clean_text(text)


def extract_txt(path):
    encs = ["utf-8", "gbk", "gb2312", "latin-1"]
    for enc in encs:
        try:
            f = open(path, "r", encoding=enc)
            text = f.read()
            f.close()
            if text and "\ufffd" not in text[:100]:
                return clean_text(text)
        except Exception:
            continue
    raise Exception("TXT编码无法识别")


def fix_scores(data):
    sections = data.get("sections", [])
    for section in sections:
        if section.get("type") == "radar_score":
            dims = section.get("dimensions", [])
            for dim in dims:
                score = dim.get("score", 0)
                if score is None:
                    score = 70
                if isinstance(score, str):
                    try:
                        score = int(float(score))
                    except:
                        score = 70
                if score <= 10:
                    score = score * 10
                if score < 50:
                    score = 50
                if score > 100:
                    score = 100
                dim["score"] = int(score)
            if dims:
                total = sum(d["score"] for d in dims) // len(dims)
                section["total_score"] = total
            else:
                section["total_score"] = 70
    return data


def ai_parse(text):
    if len(text) > 15000:
        text = text[:15000]
    prompt = get_system_prompt()
    resp = client.chat.completions.create(
        model=AI_MODEL,
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": "请解析以下报告，输出标准化JSON：\n\n" + text}
        ],
        temperature=0.05,
        max_tokens=4000,
        top_p=0.9
    )
    result = resp.choices[0].message.content.strip()
    if result.startswith("```"):
        parts = result.split("\n", 1)
        if len(parts) > 1:
            result = parts[1]
    if result.endswith("```"):
        result = result[:-3]
    result = result.strip()
    json_match = re.search(r'\{[\s\S]*\}', result)
    if json_match:
        result = json_match.group()
    data = json.loads(result)
    if "title" not in data or "sections" not in data:
        raise ValueError("JSON缺少字段")
    data = fix_scores(data)
    return data


@app.get("/")
def root():
    return {"status": "ok", "message": "ReportViz API v2.1"}


@app.get("/api/health")
def health():
    return {"status": "healthy", "ai_model": AI_MODEL}


@app.post("/api/parse")
async def parse_report(file: Optional[UploadFile] = File(None), text: Optional[str] = Form(None)):
    input_text = ""
    if file and file.filename:
        ext = file.filename.split(".")[-1].lower()
        content = await file.read()
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix="." + ext)
        tmp.write(content)
        tmp_path = tmp.name
        tmp.close()
        try:
            if ext == "pdf":
                input_text = extract_pdf(tmp_path)
            elif ext in ("docx", "doc"):
                input_text = extract_docx(tmp_path)
            elif ext == "txt":
                input_text = extract_txt(tmp_path)
            else:
                return JSONResponse(status_code=400, content={"detail": "不支持的格式"})
        except Exception as e:
            return JSONResponse(status_code=400, content={"detail": str(e)})
        finally:
            os.unlink(tmp_path)
    elif text and text.strip():
        input_text = clean_text(text.strip())
    else:
        return JSONResponse(status_code=400, content={"detail": "请上传文件或输入文本"})
    if len(input_text) < 30:
        return JSONResponse(status_code=400, content={"detail": "内容太短"})
    try:
        result = ai_parse(input_text)
        return JSONResponse(content=result)
    except json.JSONDecodeError as e:
        return JSONResponse(status_code=500, content={"detail": "AI返回格式错误: " + str(e)})
    except Exception as e:
        return JSONResponse(status_code=500, content={"detail": "解析失败: " + str(e)})


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port, workers=2, timeout_keep_alive=120)
